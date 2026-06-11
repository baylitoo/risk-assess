"""Integration tests: parse_document() → artifact chunks, intake wiring."""
from __future__ import annotations

from pathlib import Path

import pytest

from riskos.ids import stable_id
from riskos.intake import ingest_directory
from riskos.parsers import ParseConfig, parse_document
from riskos.schemas.artifacts import DocumentChunk, ImageChunk

from .pdf_fixtures import make_text_pdf


# ---------------------------------------------------------------------------
# parse_document
# ---------------------------------------------------------------------------

def test_text_pdf_produces_document_chunks(tmp_path):
    path = make_text_pdf(
        tmp_path,
        [
            "NEXUS PAYMENTS HUB Kong API Gateway TLS 1.3 internet-facing",
            "Payment Engine Java 21 Spring Boot 3.2 transaction processing",
            "PostgreSQL 16.3 customer PII encrypted at rest replication",
        ],
    )
    doc_id = stable_id("document", "asm-test", "doc.pdf")
    chunks, img_chunks = parse_document(path, doc_id)

    assert len(chunks) > 0
    assert img_chunks == []
    assert all(isinstance(c, DocumentChunk) for c in chunks)
    assert all(c.document_id == doc_id for c in chunks)
    combined = " ".join(c.text for c in chunks)
    # At least some text should be extracted
    assert len(combined) > 10


def test_chunk_ids_are_stable_across_calls(tmp_path):
    path = make_text_pdf(tmp_path, ["Stable content for ID test"])
    doc_id = stable_id("document", "asm-stable", "doc.pdf")
    chunks1, _ = parse_document(path, doc_id)
    chunks2, _ = parse_document(path, doc_id)
    ids1 = {c.chunk_id for c in chunks1}
    ids2 = {c.chunk_id for c in chunks2}
    assert ids1 == ids2


def test_docx_produces_document_chunks(tmp_path):
    from docx import Document

    doc_path = tmp_path / "spec.docx"
    doc = Document()
    doc.add_paragraph("SWIFT adapter SWIFTNet Link 7.5 critical third party")
    doc.add_paragraph("FraudGuard AI ONNX Runtime 1.18 not an LLM")
    doc.save(str(doc_path))

    doc_id = stable_id("document", "asm-docx", "spec.docx")
    chunks, img_chunks = parse_document(doc_path, doc_id)

    assert len(chunks) > 0
    assert img_chunks == []
    combined = " ".join(c.text for c in chunks)
    assert "SWIFT" in combined or "FraudGuard" in combined


def test_force_vision_path_produces_image_chunks(tmp_path):
    path = make_text_pdf(tmp_path, ["Some text that would normally be text path"])
    doc_id = stable_id("document", "asm-vision", "doc.pdf")
    cfg = ParseConfig(force_path="vision")
    chunks, img_chunks = parse_document(path, doc_id, config=cfg)

    assert chunks == []
    assert len(img_chunks) > 0
    assert all(isinstance(c, ImageChunk) for c in img_chunks)
    assert all(c.media_type == "image/jpeg" for c in img_chunks)
    assert all(c.document_id == doc_id for c in img_chunks)
    assert all(c.page_num >= 1 for c in img_chunks)
    assert all(len(c.data_b64) > 0 for c in img_chunks)
    assert all(c.sha256 for c in img_chunks)


def test_image_chunk_ids_are_stable(tmp_path):
    path = make_text_pdf(tmp_path, ["Vision content"])
    doc_id = stable_id("document", "asm-img-stable", "doc.pdf")
    cfg = ParseConfig(force_path="vision")
    _, imgs1 = parse_document(path, doc_id, config=cfg)
    _, imgs2 = parse_document(path, doc_id, config=cfg)
    assert {c.chunk_id for c in imgs1} == {c.chunk_id for c in imgs2}


def test_unsupported_extension_raises():
    with pytest.raises(ValueError, match="unsupported file type"):
        parse_document(Path("doc.xyz"), "doc-id")


# ---------------------------------------------------------------------------
# intake wiring
# ---------------------------------------------------------------------------

def test_ingest_directory_handles_pdf(tmp_path):
    make_text_pdf(
        tmp_path,
        ["API Gateway Kong 3.6 internet-facing", "Payment processor Java Spring"],
        "sats.pdf",
    )
    corpus = ingest_directory(tmp_path, stable_id("assessment", "pdf-intake"))
    assert len(corpus.documents) == 1
    assert corpus.documents[0].media_type == "application/pdf"
    assert len(corpus.chunks) > 0
    assert corpus.issues == []


def test_ingest_directory_handles_mixed_formats(tmp_path):
    make_text_pdf(tmp_path, ["PDF content about security"], "arch.pdf")
    (tmp_path / "policy.md").write_text("# Policy\n\nControl requirements.", encoding="utf-8")

    corpus = ingest_directory(tmp_path, stable_id("assessment", "mixed"))
    assert len(corpus.documents) == 2
    assert len(corpus.chunks) > 0
    assert corpus.issues == []


def test_vision_pdf_lands_in_image_chunks(tmp_path):
    make_text_pdf(tmp_path, ["Scanned page"], "scan.pdf")
    cfg = ParseConfig(force_path="vision")
    corpus = ingest_directory(
        tmp_path,
        stable_id("assessment", "vision-intake"),
        parse_config=cfg,
    )
    assert len(corpus.image_chunks) > 0
    assert corpus.chunks == []


def test_pdf_chunk_ids_reference_back_to_document(tmp_path):
    make_text_pdf(tmp_path, ["Chunk reference test content"], "ref.pdf")
    corpus = ingest_directory(tmp_path, stable_id("assessment", "ref-test"))
    doc = corpus.documents[0]
    chunk_ids_in_doc = set(doc.chunk_ids)
    chunk_ids_in_corpus = {c.chunk_id for c in corpus.chunks}
    assert chunk_ids_in_doc == chunk_ids_in_corpus

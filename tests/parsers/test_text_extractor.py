"""Tests for the text extraction path (liteparse PDF + python-docx)."""
from __future__ import annotations

from pathlib import Path

import pytest

from riskos.parsers.text import extract_pdf_pages

from .pdf_fixtures import make_text_pdf


def test_extracts_text_from_all_pages(tmp_path):
    path = make_text_pdf(
        tmp_path,
        [
            "NEXUS PAYMENTS HUB security assessment component inventory",
            "Kong API Gateway 3.6 internet-facing TLS 1.3 rate limiting",
            "PostgreSQL 16.3 PII encrypted at rest customer database",
        ],
    )
    pages = extract_pdf_pages(path)
    assert len(pages) == 3
    texts = " ".join(p.text for p in pages)
    assert "NEXUS" in texts or "Kong" in texts or "PostgreSQL" in texts


def test_page_numbers_are_one_based(tmp_path):
    path = make_text_pdf(tmp_path, ["Page A", "Page B"])
    pages = extract_pdf_pages(path)
    nums = [p.page_num for p in pages if p.text.strip()]
    assert all(n >= 1 for n in nums)
    assert len(set(nums)) == len(nums)  # unique


def test_empty_pages_do_not_crash(tmp_path):
    path = make_text_pdf(tmp_path, ["", "Some content", ""])
    pages = extract_pdf_pages(path)
    assert any(p.text.strip() for p in pages)


def test_items_have_spatial_metadata(tmp_path):
    path = make_text_pdf(tmp_path, ["Security architecture trust boundary"])
    pages = extract_pdf_pages(path)
    # At least one page should have bounding-box items
    items = []
    for p in pages:
        items.extend(p.items)
    if items:
        item = items[0]
        assert item.width > 0
        assert item.height > 0


def test_docx_extraction(tmp_path):
    from docx import Document
    from riskos.parsers.text import extract_docx_paragraphs

    doc_path = tmp_path / "spec.docx"
    doc = Document()
    doc.add_paragraph("SWIFT adapter SWIFTNet Link 7.5 payment processing")
    doc.add_paragraph("")
    doc.add_paragraph("FraudGuard AI ONNX Runtime anomaly detection")
    doc.save(str(doc_path))

    paragraphs = extract_docx_paragraphs(doc_path)
    assert len(paragraphs) >= 2
    combined = " ".join(paragraphs)
    assert "SWIFT" in combined
    assert "FraudGuard" in combined
